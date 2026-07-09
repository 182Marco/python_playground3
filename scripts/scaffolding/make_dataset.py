"""Genera il dataset "sporco" del lab in dataset/ — 9 file, 9 situazioni reali.

Ogni file è costruito per rompersi (o sporcarsi) in un modo DIVERSO, e per
SEMBRARE VERO: multipagina, header/footer ripetuti, boilerplate abbondante.

  guida_smartworking.pdf   → policy di 3 pagine, header ripetuto + numeri di
                             pagina (il caso "pulito"... ma il cleaning lavora)
  listino_prezzi.pdf       → 2 pagine con TABELLA di ~28 righe: pypdf estrae
                             le celle alla rinfusa
  verbale_scansionato.pdf  → 2 pagine "scansione": solo immagini (con grana e
                             pagina storta), zero layer di testo
  report_trimestrale.pdf   → report troncato al 55%: il loader va in eccezione
  faq_prodotto.html        → pagina e-commerce vera: cookie banner, menu,
                             breadcrumb, promo, sidebar, recensioni, newsletter,
                             footer a colonne — tutto in <div> (il loader L16
                             toglie <nav>/<footer>, qui non basta)
  contratto_fornitura.docx → contratto di 8 articoli con TABELLE (SLA e
                             corrispettivi): il loader legge solo i paragrafi,
                             le tabelle spariscono IN SILENZIO
  note_interne.txt         → appunti interni in latin-1: UnicodeDecodeError

    python make_dataset.py

Serve solo per (ri)generare il dataset (già versionato). Richiede `reportlab` e
`pillow` (vedi requirements.txt, sezione "solo rigenerazione").
"""
import random
from pathlib import Path

DATASET = Path(__file__).parent.parent / "dataset"


# ================================================================= contenuti
GUIDA_SEZIONI = [
    ("Premessa e ambito di applicazione",
     "Il presente regolamento disciplina le modalità di svolgimento della "
     "prestazione lavorativa in modalità agile (di seguito, smart working) per "
     "tutto il personale dipendente di Lumen S.r.l., a decorrere dal 1 marzo "
     "2026. Lo smart working è una modalità di esecuzione del rapporto di "
     "lavoro subordinato, stabilita mediante accordo individuale tra le parti, "
     "senza precisi vincoli di orario o di luogo di lavoro, nel rispetto della "
     "normativa vigente e degli accordi collettivi applicabili. Il presente "
     "documento annulla e sostituisce ogni precedente comunicazione in materia."),
    ("Definizioni",
     "Ai fini del presente regolamento si intende per: giornata agile, la "
     "giornata lavorativa svolta al di fuori dei locali aziendali; accordo "
     "individuale, l'accordo scritto tra dipendente e azienda che disciplina "
     "l'adesione allo smart working; fascia di reperibilità, l'intervallo "
     "orario nel quale il dipendente garantisce la propria raggiungibilità; "
     "dotazione aziendale, l'insieme degli strumenti informatici forniti da "
     "Lumen S.r.l. per lo svolgimento della prestazione."),
    ("Giornate di lavoro agile",
     "Ogni dipendente può svolgere fino a tre giornate settimanali in modalità "
     "agile, da concordare con il proprio responsabile entro il venerdì della "
     "settimana precedente tramite il gestionale presenze. Le giornate non "
     "godute non sono cumulabili nelle settimane successive, né monetizzabili. "
     "Il responsabile può richiedere la presenza in sede per esigenze "
     "organizzative motivate, con preavviso di almeno ventiquattro ore. Nei "
     "periodi di chiusura collettiva e durante il periodo di prova lo smart "
     "working è sospeso, salvo diversa autorizzazione della Direzione."),
    ("Richiesta, approvazione e revoca",
     "L'adesione allo smart working avviene su base volontaria, mediante "
     "sottoscrizione dell'accordo individuale. La richiesta va presentata "
     "all'ufficio HR con almeno quindici giorni di anticipo. L'azienda può "
     "revocare l'accordo con preavviso di trenta giorni, ovvero con effetto "
     "immediato in presenza di giustificato motivo. Il dipendente può recedere "
     "dall'accordo con preavviso di quindici giorni."),
    ("Strumenti di lavoro e sicurezza informatica",
     "L'azienda fornisce laptop aziendale, accesso VPN e licenze software "
     "necessarie. È vietato l'uso di reti Wi-Fi pubbliche non protette per "
     "l'accesso ai sistemi aziendali; è altresì vietato l'utilizzo di "
     "dispositivi personali per il trattamento di dati riservati, salvo "
     "autorizzazione scritta del responsabile IT. Le credenziali di accesso "
     "sono strettamente personali e non cedibili. Ogni anomalia di sicurezza "
     "va segnalata entro quattro ore all'indirizzo security@lumen.example."),
    ("Reperibilità e diritto alla disconnessione",
     "Durante le giornate di lavoro agile il dipendente garantisce la "
     "reperibilità nelle fasce orarie 9:30-12:30 e 14:30-17:00. Al di fuori "
     "di tali fasce trova applicazione il diritto alla disconnessione: il "
     "dipendente non è tenuto a leggere o rispondere a comunicazioni "
     "lavorative, e l'eventuale mancata risposta non può costituire oggetto "
     "di valutazione disciplinare. Le riunioni vanno pianificate di norma "
     "all'interno delle fasce di reperibilità."),
    ("Salute e sicurezza",
     "Il datore di lavoro consegna al dipendente e al rappresentante dei "
     "lavoratori per la sicurezza un'informativa scritta sui rischi generali "
     "e specifici connessi alla modalità agile. Il dipendente è tenuto a "
     "cooperare all'attuazione delle misure di prevenzione, scegliendo "
     "postazioni di lavoro conformi ai principi di ergonomia e adeguatamente "
     "illuminate. L'INAIL copre gli infortuni occorsi durante la prestazione "
     "agile secondo la normativa vigente."),
    ("Rimborsi e dotazioni accessorie",
     "Non sono previsti rimborsi forfettari per le utenze domestiche. "
     "Eventuali acquisti di dotazioni ergonomiche (sedute, supporti, monitor "
     "aggiuntivi) vanno autorizzati preventivamente dall'ufficio acquisti "
     "tramite richiesta motivata; in caso di approvazione, il bene resta di "
     "proprietà aziendale e va restituito alla cessazione del rapporto. La "
     "stampa di documenti in modalità agile è di norma esclusa."),
    ("Luoghi di lavoro e trasferte",
     "La prestazione agile può essere svolta in qualunque luogo idoneo sul "
     "territorio nazionale che garantisca condizioni di sicurezza, "
     "riservatezza e connettività adeguata. È esclusa la prestazione da "
     "luoghi pubblici affollati per attività che comportino il trattamento "
     "di dati riservati. Lo svolgimento della prestazione dall'estero deve "
     "essere autorizzato preventivamente dall'ufficio HR per le necessarie "
     "verifiche fiscali e previdenziali. Le giornate di trasferta non sono "
     "computate come giornate agili."),
    ("Comportamenti e riservatezza",
     "Durante le videochiamate il dipendente cura che l'ambiente inquadrato "
     "sia consono al contesto professionale. I documenti cartacei contenenti "
     "dati aziendali vanno custoditi con diligenza e distrutti in modo "
     "sicuro al termine dell'utilizzo. È vietato consentire a terzi, inclusi "
     "i familiari, l'utilizzo della dotazione aziendale. Le riunioni con "
     "contenuti riservati vanno svolte con auricolari."),
    ("Formazione e monitoraggio",
     "L'azienda organizza sessioni formative periodiche su sicurezza "
     "informatica, strumenti di collaborazione e gestione del tempo. La "
     "prestazione in modalità agile è valutata per obiettivi e risultati, con "
     "gli stessi criteri applicati alla prestazione in sede. È escluso ogni "
     "controllo a distanza non conforme alla normativa sul lavoro."),
    ("Decorrenza, durata e revisione",
     "Il presente regolamento entra in vigore il 1 marzo 2026 e ha durata "
     "indeterminata. La Direzione si riserva di rivederlo con cadenza annuale, "
     "previo confronto con le rappresentanze dei lavoratori. Per quanto non "
     "espressamente previsto si rinvia alla legge 81/2017 e al CCNL applicato. "
     "Per chiarimenti: hr@lumen.example, interno 214."),
]

LISTINO_INTRO = (
    "Listino prezzi 2026 in vigore dal 1 febbraio 2026. Prezzi in euro, IVA "
    "esclusa, franco nostro magazzino di Milano. Il presente listino annulla e "
    "sostituisce ogni versione precedente. Per ordini superiori a 50 unità "
    "dello stesso articolo si applica lo sconto volume indicato in tabella; "
    "gli sconti non sono cumulabili con altre promozioni in corso.")

LISTINO_RIGHE = [
    ["Codice", "Prodotto", "Prezzo unitario", "Sconto volume"],
    ["LD-100", "LumaDesk Basic 120x70", "129,00", "5%"],
    ["LD-110", "LumaDesk Basic 140x70", "149,00", "5%"],
    ["LD-120", "LumaDesk Basic 160x80", "169,00", "5%"],
    ["LD-200", "LumaDesk Pro 140x70", "249,00", "8%"],
    ["LD-210", "LumaDesk Pro 160x80", "279,00", "8%"],
    ["LD-220", "LumaDesk Pro 180x80", "309,00", "8%"],
    ["LD-300", "LumaDesk Enterprise 160x80", "499,00", "12%"],
    ["LD-310", "LumaDesk Enterprise 180x90", "549,00", "12%"],
    ["LD-320", "LumaDesk Enterprise 200x90", "599,00", "12%"],
    ["LD-400", "LumaDesk Conference 240x110", "899,00", "15%"],
    ["LL-010", "LumaLight lampada da scrivania", "59,00", "5%"],
    ["LL-020", "LumaLight Pro con sensore", "89,00", "5%"],
    ["LL-030", "LumaBar barra monitor", "74,00", "5%"],
    ["LK-050", "Kit accessori (canalina + 3 passacavi)", "39,00", "—"],
    ["LK-060", "Cassettiera con serratura", "119,00", "5%"],
    ["LK-070", "Supporto CPU regolabile", "45,00", "—"],
    ["LK-080", "Braccio monitor singolo", "79,00", "5%"],
    ["LK-081", "Braccio monitor doppio", "129,00", "5%"],
    ["LK-090", "Tappetino antifatica", "35,00", "—"],
    ["LM-100", "Piano di ricambio bambù 140x70", "95,00", "—"],
    ["LM-110", "Piano di ricambio bambù 160x80", "115,00", "—"],
    ["LM-200", "Centralina di ricambio", "65,00", "—"],
    ["SRV-01", "Installazione on-site (a postazione)", "180,00", "n/a"],
    ["SRV-02", "Installazione on-site oltre 10 postazioni", "150,00", "n/a"],
    ["SRV-03", "Estensione garanzia +2 anni", "49,00", "n/a"],
    ["SRV-04", "Manutenzione programmata annuale", "120,00", "n/a"],
    ["SRV-05", "Ritiro e smaltimento usato", "60,00", "n/a"],
]

LISTINO_CONDIZIONI = [
    ("Consegna", "Tempi di consegna indicativi: 5 giorni lavorativi per gli "
     "articoli a magazzino, 15 giorni per le configurazioni Enterprise e "
     "Conference. La consegna al piano è inclusa per ordini superiori a "
     "1.500 euro."),
    ("Garanzia", "Tutti i prodotti LumaDesk godono di garanzia di 5 anni sul "
     "motore e 2 anni sulle parti elettroniche. L'estensione SRV-03 porta la "
     "copertura delle parti elettroniche a 4 anni."),
    ("Resi", "Eventuali resi vanno autorizzati dal servizio clienti entro 14 "
     "giorni dalla consegna; la merce deve essere integra e nell'imballo "
     "originale. Le spese di reso sono a carico del cliente salvo prodotto "
     "difettoso."),
    ("Validità", "Il presente listino è valido fino al 31 gennaio 2027, salvo "
     "esaurimento scorte o variazioni straordinarie dei costi delle materie "
     "prime, che saranno comunicate con 30 giorni di preavviso."),
]

VERBALE_PAG1 = """LUMEN S.R.L. - VERBALE DI RIUNIONE N. 14/2026

Data: 12 giugno 2026, ore 14:30 - Sala riunioni, sede di Milano

Presenti: M. Rota (direzione generale), G. Bianchi (IT),
S. Conti (ufficio acquisti), L. Ferri (HR), P. Galli (amministrazione).
Assenti giustificati: R. Marino (commerciale).

Ordine del giorno:
1. Migrazione del gestionale in cloud
2. Budget formazione secondo semestre
3. Rinnovo contratto di pulizie
4. Varie ed eventuali

Punto 1 - Migrazione del gestionale.
Bianchi illustra le tre offerte ricevute (NuvolaTech, CloudItalia,
DataCenter Sud). Dopo ampia discussione, la migrazione e' approvata
all'unanimita': avvio il 1 settembre 2026, fornitore NuvolaTech S.p.A.,
budget massimo 85.000 euro nel triennio. Conti predisporra' il contratto
di fornitura entro fine giugno; Bianchi il piano di migrazione con
particolare attenzione alla qualita' dei dati esistenti."""

VERBALE_PAG2 = """Punto 2 - Budget formazione.
Ferri presenta il consuntivo del primo semestre (18.500 euro) e propone
un aumento per il secondo semestre, motivato dal progetto cloud. Il
budget formazione e' portato a 40.000 euro, con priorita' ai corsi su
sicurezza informatica e strumenti di collaborazione.

Punto 3 - Rinnovo contratto di pulizie.
Rinviato: Conti deve ancora ricevere due preventivi. Se ne riparla
alla prossima riunione.

Punto 4 - Varie ed eventuali.
Galli segnala ritardi nei rimborsi spese di maggio, in smaltimento
entro il 20 giugno. Rota ricorda la chiusura estiva (11-22 agosto).

Il verbale e' approvato all'unanimita'.
Prossima riunione: 10 luglio 2026, ore 15:00.

Il segretario verbalizzante          Il presidente
     P. Galli                            M. Rota"""

REPORT_PARAGRAFI = [
    "Nel secondo trimestre 2026 il fatturato consolidato del gruppo Lumen ha "
    "raggiunto 12,4 milioni di euro, in crescita del 9,5 per cento rispetto "
    "allo stesso periodo dell'esercizio precedente. La crescita è trainata "
    "dalla linea LumaDesk Pro e dal canale e-commerce, che rappresenta ormai "
    "il 34 per cento delle vendite complessive.",
    "Il margine operativo lordo si attesta al 17,8 per cento, in lieve "
    "flessione rispetto al primo trimestre per effetto dell'aumento dei costi "
    "di trasporto e del bambù certificato. La direzione acquisti ha avviato "
    "la rinegoziazione dei contratti di fornitura con effetto dal quarto "
    "trimestre.",
    "Sul fronte operativo, il progetto di migrazione del gestionale in cloud "
    "procede secondo il cronoprogramma approvato: la fase di assessment della "
    "qualità dei dati si è conclusa a giugno, evidenziando la necessità di un "
    "intervento straordinario di bonifica degli archivi documentali.",
    "Le previsioni per il terzo trimestre restano prudenti in ragione della "
    "stagionalità del mercato office. Il portafoglio ordini al 30 giugno "
    "copre il 61 per cento del budget trimestrale.",
]

FAQ_HTML = """<!DOCTYPE html>
<html lang="it">
<head><meta charset="utf-8"><title>FAQ LumaDesk — domande frequenti | Lumen Store</title></head>
<body>
<div class="cookie-banner">
  Questo sito utilizza cookie tecnici e, previo tuo consenso, cookie di
  profilazione nostri e di terze parti per mostrarti pubblicità in linea con
  le tue preferenze. Cliccando su "Accetta tutti" acconsenti all'uso dei
  cookie. Puoi modificare le tue scelte in ogni momento dalla pagina
  Preferenze cookie.
  <a href="#">Accetta tutti</a> <a href="#">Solo necessari</a> <a href="#">Preferenze</a>
</div>
<div class="topbar">
  <a href="/">Home</a> | <a href="/prodotti">Prodotti</a> |
  <a href="/listino">Listino</a> | <a href="/assistenza">Assistenza</a> |
  <a href="/azienda">Azienda</a> | <a href="/contatti">Contatti</a> |
  <a href="/account">Il mio account</a> | <a href="/carrello">Carrello (0)</a>
</div>
<div class="promo-strip">
  SUMMER SALE: fino al 31 luglio -15% su tutta la gamma LumaDesk con il codice
  ESTATE15. Spedizione gratuita sopra i 199 euro!
</div>
<div class="breadcrumb">Home › Prodotti › LumaDesk › Domande frequenti</div>
<div class="content">
  <h1>Domande frequenti su LumaDesk</h1>

  <h2>Che cos'è LumaDesk?</h2>
  <p>LumaDesk è la scrivania regolabile in altezza di Lumen S.r.l., pensata
  per il lavoro da casa e per l'ufficio. È disponibile nelle versioni Basic,
  Pro, Enterprise e Conference, con piani da 120x70 a 240x110 cm. Tutte le
  versioni hanno il piano in bambù certificato FSC e il doppio motore
  silenzioso sotto i 40 dB.</p>

  <h2>Quali sono le differenze tra Basic, Pro ed Enterprise?</h2>
  <p>La versione Basic ha regolazione a due segmenti (71-119 cm) e memoria
  singola. La Pro aggiunge il terzo segmento (60-125 cm), quattro memorie di
  posizione e la ricarica USB-C integrata. La Enterprise include anche il
  sensore anticollisione, il modulo di prenotazione postazione e la
  centralina compatibile con il software di gestione flotte.</p>

  <h2>Quanto dura la garanzia?</h2>
  <p>La garanzia è di 5 anni sul motore e di 2 anni sulle parti elettroniche.
  Per i clienti Enterprise la garanzia sul motore è estesa a 8 anni. Con il
  servizio SRV-03 puoi estendere di ulteriori 2 anni la copertura delle parti
  elettroniche.</p>

  <h2>Di che materiale è il piano?</h2>
  <p>Il piano è in bambù massello certificato, trattato con olio naturale
  atossico. Su richiesta sono disponibili piani in laminato bianco o rovere.
  Il bambù è un materiale vivo: piccole variazioni di venatura e colore sono
  normali e non costituiscono difetto.</p>

  <h2>Quanto peso sopporta?</h2>
  <p>La portata dinamica è di 80 kg per Basic e Pro e di 120 kg per
  Enterprise e Conference, distribuiti uniformemente sul piano. Nel calcolo
  vanno considerati monitor, bracci, cassettiere agganciate e l'eventuale
  supporto CPU.</p>

  <h2>Come richiedo assistenza?</h2>
  <p>Puoi aprire un ticket dall'area clienti oppure scrivere a
  assistenza@lumen.example indicando il numero di serie (etichetta sotto il
  piano, lato sinistro). Il tempo di prima risposta è di 1 giorno lavorativo;
  per i clienti Enterprise con contratto attivo è di 4 ore lavorative.</p>

  <h2>Posso montarla da solo?</h2>
  <p>Sì: il montaggio richiede circa 30 minuti, un cacciavite a croce e la
  chiave a brugola inclusa nella confezione. Il manuale illustrato è nella
  scatola e in versione video nell'area clienti. Per ordini aziendali è
  disponibile l'installazione on-site (codice SRV-01 a listino).</p>

  <h2>In quanto tempo arriva?</h2>
  <p>Gli articoli a magazzino vengono spediti entro 2 giorni lavorativi con
  corriere dedicato; la consegna avviene in 3-5 giorni. Le versioni
  Enterprise e Conference sono prodotte su ordinazione: tempi indicativi 15
  giorni lavorativi.</p>

  <h2>Come funzionano i resi?</h2>
  <p>Hai 14 giorni dalla consegna per richiedere il reso dall'area clienti.
  La scrivania deve essere smontata e riposta nell'imballo originale; il
  ritiro viene organizzato dal nostro corriere. Il rimborso avviene entro 14
  giorni dal rientro della merce in magazzino.</p>

  <h2>La app è compatibile con la mia scrivania?</h2>
  <p>La app Lumen Move (iOS e Android) funziona con Pro, Enterprise e
  Conference prodotte dal 2024 in poi. Ti ricorda di alternare le posizioni,
  memorizza le altezze preferite e — sulle Enterprise — si integra con la
  lampada LumaLight Pro per regolare l'illuminazione della postazione.</p>
</div>
<div class="sidebar">
  Ti potrebbero interessare
  <a href="#">LumaLight Pro con sensore — aggiungi al carrello</a>
  <a href="#">Braccio monitor doppio — aggiungi al carrello</a>
  <a href="#">Cassettiera con serratura — aggiungi al carrello</a>
  <a href="#">Tappetino antifatica — aggiungi al carrello</a>
</div>
<div class="reviews-teaser">
  ★★★★★ 4,8/5 su 1.243 recensioni verificate — leggi tutte le recensioni dei
  nostri clienti
</div>
<div class="newsletter">
  Iscriviti alla newsletter di Lumen per ricevere offerte esclusive, consigli
  di ergonomia e anteprime sui nuovi prodotti! Inserisci la tua email e
  accetta la privacy policy per iniziare.
</div>
<div class="footer-links">
  Chi siamo | Lavora con noi | Sostenibilità | Rivenditori autorizzati |
  Spedizioni e resi | Privacy policy | Termini e condizioni | Cookie policy |
  © 2026 Lumen S.r.l. — P.IVA 01234567890 — Cap. soc. 100.000 euro i.v. |
  Seguici su
  Facebook Instagram LinkedIn YouTube
</div>
</body>
</html>
"""

CONTRATTO_PAR = [
    ("h1", "Contratto di fornitura di servizi cloud"),
    ("p", "Tra Lumen S.r.l., con sede legale in Milano, via dei Tigli 14, "
          "P.IVA 01234567890, in persona del direttore generale M. Rota "
          "(di seguito, il Cliente), e NuvolaTech S.p.A., con sede legale in "
          "Torino, corso Stati Uniti 8, P.IVA 09876543210, in persona del "
          "procuratore speciale A. Ricciardi (di seguito, il Fornitore), si "
          "conviene e si stipula quanto segue."),
    ("h2", "Premesse"),
    ("p", "Il Cliente ha deliberato, con verbale di riunione n. 14/2026 del "
          "12 giugno 2026, la migrazione del proprio gestionale aziendale in "
          "cloud. Il Fornitore è operatore specializzato in servizi di hosting "
          "gestito e migrazione applicativa. Le premesse costituiscono parte "
          "integrante e sostanziale del presente contratto."),
    ("h2", "Art. 1 — Oggetto"),
    ("p", "Il Fornitore si impegna a erogare in favore del Cliente: (a) i "
          "servizi di migrazione del gestionale e degli archivi documentali; "
          "(b) i servizi di hosting gestito in alta affidabilità; (c) i "
          "servizi di backup e disaster recovery; (d) il supporto tecnico "
          "specialistico, secondo i livelli di servizio di cui all'art. 4."),
    ("h2", "Art. 2 — Durata"),
    ("p", "Il contratto ha durata di 36 mesi a decorrere dal 1 settembre 2026 "
          "e si intende rinnovato per ulteriori 12 mesi salvo disdetta di una "
          "delle parti con preavviso di 90 giorni, da comunicarsi a mezzo PEC."),
    ("h2", "Art. 3 — Corrispettivi e fatturazione"),
    ("p", "I corrispettivi dovuti dal Cliente sono riportati nella tabella "
          "seguente, parte integrante del presente contratto. La fatturazione "
          "è trimestrale anticipata, con pagamento a 60 giorni data fattura."),
    ("tabella", "corrispettivi"),
    ("h2", "Art. 4 — Livelli di servizio (SLA)"),
    ("p", "Il Fornitore garantisce i livelli di servizio riportati nella "
          "tabella seguente. Il mancato rispetto degli SLA comporta "
          "l'applicazione delle penali indicate, che il Cliente potrà "
          "compensare con i corrispettivi dovuti."),
    ("tabella", "sla"),
    ("h2", "Art. 5 — Obblighi del Fornitore"),
    ("p", "Il Fornitore si impegna a: mantenere le certificazioni ISO 27001 e "
          "ISO 9001 per tutta la durata contrattuale; localizzare i dati del "
          "Cliente esclusivamente in data center situati nell'Unione Europea; "
          "comunicare ogni incidente di sicurezza entro 24 ore dalla "
          "rilevazione; fornire report mensili sui livelli di servizio."),
    ("h2", "Art. 6 — Riservatezza e protezione dei dati"),
    ("p", "Le parti si impegnano a mantenere riservate le informazioni "
          "scambiate in esecuzione del presente contratto. Il Fornitore è "
          "nominato responsabile del trattamento ai sensi dell'art. 28 del "
          "Regolamento UE 2016/679, secondo l'accordo allegato sub B."),
    ("h2", "Art. 7 — Recesso e risoluzione"),
    ("p", "Il Cliente può recedere in qualsiasi momento con preavviso di 90 "
          "giorni, corrispondendo i soli servizi erogati. Il contratto si "
          "risolve di diritto in caso di violazione degli obblighi di cui "
          "agli artt. 5 e 6, previa diffida ad adempiere entro 15 giorni."),
    ("h2", "Art. 8 — Foro competente"),
    ("p", "Per ogni controversia derivante dal presente contratto è "
          "competente in via esclusiva il Foro di Milano."),
    ("p", "Milano, 30 giugno 2026"),
    ("p", "Lumen S.r.l. — M. Rota          NuvolaTech S.p.A. — A. Ricciardi"),
]

CONTRATTO_SLA = [
    ["Servizio", "SLA garantito", "Penale per violazione"],
    ["Disponibilità piattaforma", "99,9% su base mensile", "5% del canone mensile"],
    ["Tempo di ripristino (RTO)", "4 ore", "250 euro per ora di ritardo"],
    ["Perdita dati massima (RPO)", "15 minuti", "1.000 euro per evento"],
    ["Prima risposta al ticket (P1)", "30 minuti, h24", "100 euro per violazione"],
    ["Prima risposta al ticket (P2)", "4 ore lavorative", "50 euro per violazione"],
    ["Report mensile SLA", "entro il giorno 5", "50 euro per giorno di ritardo"],
]

CONTRATTO_CORRISPETTIVI = [
    ["Voce", "Importo", "Periodicità"],
    ["Migrazione una tantum", "18.000 euro", "alla consegna"],
    ["Hosting gestito", "2.100 euro", "mensile"],
    ["Backup e disaster recovery", "450 euro", "mensile"],
    ["Supporto specialistico (10 h/mese)", "900 euro", "mensile"],
    ["Ore di supporto extra", "95 euro/ora", "a consumo"],
]

FATTURA_RIGHE = [
    ("COD.", "DESCRIZIONE", "Q.TA'", "PREZZO", "TOTALE"),
    ("MIG-01", "Assessment qualita' dati", "1", "4.500,00", "4.500,00"),
    ("MIG-02", "Setup ambiente cloud staging", "1", "3.200,00", "3.200,00"),
    ("SUP-10", "Supporto specialistico (ore)", "12", "95,00", "1.140,00"),
    ("FRM-05", "Formazione team IT (giornate)", "2", "780,00", "1.560,00"),
    ("SPE-01", "Spese di trasferta", "1", "240,00", "240,00"),
]

RASSEGNA_COL1 = """La scrivania che cambia
altezza e' diventata il
simbolo del nuovo ufficio
ibrido. Nel distretto del
mobile brianzolo, Lumen
S.r.l. ha chiuso il 2025
con ricavi in crescita a
doppia cifra e punta ora
ai mercati esteri.

"Il lavoro da casa non e'
una parentesi", spiega la
direzione dell'azienda,
"e chi arreda l'ufficio
domestico chiede prodotti
che durino dieci anni".

La gamma LumaDesk, nata
nel 2022, copre oggi
quattro fasce di prezzo
e rappresenta il 68 per
cento del fatturato."""

RASSEGNA_COL2 = """Il piano industriale
prevede investimenti per
2,5 milioni in tre anni:
nuova linea di montaggio,
un magazzino automatico e
la migrazione dei sistemi
informativi in cloud.

Sul fronte del prodotto,
l'azienda scommette sul
bambu' certificato e su
una filiera corta: "meno
di 400 chilometri tra il
fornitore del piano e la
linea di assemblaggio".

Gli analisti del settore
osservano pero' che la
concorrenza asiatica sta
comprimendo i margini su
tutta la fascia entry
level del mercato."""

NOTE_INTERNE = """NOTE INTERNE - ufficio acquisti (bozza, NON distribuire)
aggiornate al 30/06 da S.C.

*** PRIORITÀ DI LUGLIO ***

1) MIGRAZIONE GESTIONALE - verificare la qualità dei dati PRIMA della
   migrazione di settembre. Bianchi dice che dall'assessment escono fuori
   parecchi problemi: anagrafiche doppie, allegati mancanti, encoding
   ballerino nei file più vecchi (roba del 2015-2018).
   -> chiedere a NuvolaTech il piano di ROLLBACK dettagliato (non è nel
      contratto, farlo aggiungere come allegato!!)
   -> l'inventario è da rifare: più di metà delle schede è incompleta

2) CONTRATTO NUVOLATECH - bozza ricevuta il 28/06, girata a legale.
   Occhio all'art. 3: la fatturazione anticipata non ci era stata detta
   a voce. Sentire Ricciardi (cell. finisce con 882, chiedere a Rota).

3) LISTINO 2027 - i fornitori del bambù hanno preannunciato rincari del
   7-9%. Se confermato, il listino 2027 va rivisto a ottobre e non a
   gennaio. Parlarne al prossimo direttivo.

4) VARIE
   - rimborsi maggio: SBLOCCATI il 19/06 (Galli)
   - preventivi pulizie: ne mancano ancora DUE, sollecitare
   - stampante piano 2: in garanzia, il tecnico passa lunedì

Nota di M. Rota (a voce, 27/06): la società è già in ritardo sul
cronoprogramma; se entro venerdì non c'è l'ok della direzione sul
budget rivisto, l'attività slitta a settembre. URGENTISSIMO.
"""


# ================================================================ generatori
def _stili():
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.styles import getSampleStyleSheet

    styles = getSampleStyleSheet()
    styles["BodyText"].alignment = TA_JUSTIFY
    styles["BodyText"].fontSize = 10.5
    styles["BodyText"].leading = 15.5
    return styles


def _header_footer(titolo: str):
    """Header ripetuto + numero di pagina: il noise 'vero' dei PDF aziendali."""
    from reportlab.lib.pagesizes import A4

    def draw(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.38, 0.33, 0.55)
        canvas.drawString(20 * 2.834, A4[1] - 12 * 2.834, titolo)
        canvas.drawRightString(A4[0] - 20 * 2.834, A4[1] - 12 * 2.834,
                               "Uso interno - Rev. 2026-02")
        canvas.setFillColorRGB(0.3, 0.3, 0.3)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(A4[0] / 2, 10 * 2.834, str(canvas.getPageNumber()))
        canvas.restoreState()
    return draw


def make_guida(nome: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = _stili()
    flow = [Paragraph("Regolamento per il lavoro agile (smart working)",
                      styles["Title"]),
            Paragraph("Lumen S.r.l. — Direzione Risorse Umane — edizione "
                      "febbraio 2026", styles["Italic"]),
            Spacer(1, 14)]
    for i, (titolo, testo) in enumerate(GUIDA_SEZIONI, 1):
        flow.append(Paragraph(f"{i}. {titolo}", styles["Heading2"]))
        flow.append(Paragraph(testo, styles["BodyText"]))
        flow.append(Spacer(1, 12))
    hf = _header_footer("Lumen S.r.l. - Regolamento smart working")
    SimpleDocTemplate(str(DATASET / nome), pagesize=A4,
                      topMargin=60, bottomMargin=55).build(
        flow, onFirstPage=hf, onLaterPages=hf)
    print(f"  scritto {nome} (multipagina, header/footer ripetuti)")


def make_listino(nome: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = _stili()
    tabella = Table(LISTINO_RIGHE, hAlign="LEFT", repeatRows=1,
                    colWidths=[55, 245, 90, 85])
    tabella.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a2640")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b9b3d1")),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f2effc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    flow = [Paragraph("Listino prezzi 2026", styles["Title"]),
            Paragraph(LISTINO_INTRO, styles["BodyText"]),
            Spacer(1, 12), tabella, Spacer(1, 16),
            Paragraph("Condizioni di vendita", styles["Heading2"])]
    for titolo, testo in LISTINO_CONDIZIONI:
        flow.append(Paragraph(f"<b>{titolo}.</b> {testo}", styles["BodyText"]))
        flow.append(Spacer(1, 5))
    hf = _header_footer("Lumen S.r.l. - Listino 2026")
    SimpleDocTemplate(str(DATASET / nome), pagesize=A4,
                      topMargin=60, bottomMargin=55).build(
        flow, onFirstPage=hf, onLaterPages=hf)
    print(f"  scritto {nome} (tabella di {len(LISTINO_RIGHE) - 1} righe su 2 pagine)")


def _font_mono(size: int):
    from PIL import ImageFont
    try:
        return ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", size)
    except OSError:
        return ImageFont.load_default()


def _font_bold(size: int):
    from PIL import ImageFont
    try:
        return ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", size)
    except OSError:
        return ImageFont.load_default()


def _effetto_scansione(img, seed: int, carta: str = "#f3eedd"):
    """Grana, macchie e pagina leggermente storta: il tocco 'da scanner'."""
    from PIL import ImageDraw

    rng = random.Random(seed)
    draw = ImageDraw.Draw(img)
    w, h = img.size
    for _ in range(2600):
        x, y = rng.randrange(w), rng.randrange(h)
        draw.point((x, y), fill=rng.choice(("#d8d2bd", "#c9c2ab", "#e6e0cc")))
    for _ in range(12):
        x, y = rng.randrange(w), rng.randrange(h)
        draw.ellipse((x, y, x + rng.randrange(2, 5), y + rng.randrange(2, 5)),
                     fill="#8d8570")
    return img.rotate(rng.uniform(-0.8, 0.8), expand=False, fillcolor=carta)


def _pagina_scansionata(testo: str, seed: int):
    """Una pagina di testo 'da scanner'."""
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1240, 1754), "#f3eedd")
    draw = ImageDraw.Draw(img)
    draw.multiline_text((100, 110), testo, fill="#26221c",
                        font=_font_mono(26), spacing=11)
    return _effetto_scansione(img, seed)


def _fattura_img(seed: int):
    """Una fattura scansionata CON TABELLA: OCR + struttura insieme.

    È il file 'a scala di tier': il fast la legge ma spappola la tabella,
    i tier alti la ricostruiscono.
    """
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1240, 1754), "#f6f2e6")
    d = ImageDraw.Draw(img)
    mono, mono_s, bold = _font_mono(26), _font_mono(22), _font_bold(34)

    d.text((100, 90), "NUVOLATECH S.p.A.", font=bold, fill="#1a1712")
    d.multiline_text((100, 140),
                     "Corso Stati Uniti 8, 10128 Torino\n"
                     "P.IVA 09876543210 - nuvolatech.example",
                     font=mono_s, fill="#3a352c", spacing=8)
    d.text((100, 250), "FATTURA n. 214/2026 del 30/06/2026",
           font=_font_bold(28), fill="#1a1712")
    d.multiline_text((100, 310),
                     "Spett.le Lumen S.r.l.\n"
                     "Via dei Tigli 14, 20133 Milano\n"
                     "P.IVA 01234567890",
                     font=mono_s, fill="#3a352c", spacing=8)

    # tabella righe: griglia disegnata + testo nelle celle
    x0, y0, riga_h = 100, 470, 62
    colonne = [x0, x0 + 150, x0 + 640, x0 + 750, x0 + 900, x0 + 1040]
    n = len(FATTURA_RIGHE)
    for i in range(n + 1):
        d.line((x0, y0 + i * riga_h, colonne[-1], y0 + i * riga_h),
               fill="#4a443a", width=2)
    for cx in colonne:
        d.line((cx, y0, cx, y0 + n * riga_h), fill="#4a443a", width=2)
    for i, riga in enumerate(FATTURA_RIGHE):
        font = mono_s if i else _font_mono(22)
        for testo, cx in zip(riga, colonne):
            d.text((cx + 12, y0 + i * riga_h + 18), testo, font=font,
                   fill="#26221c")

    y = y0 + n * riga_h + 50
    d.multiline_text((720, y),
                     "Imponibile      10.640,00\n"
                     "IVA 22%          2.340,80\n"
                     "TOTALE          12.980,80",
                     font=mono, fill="#1a1712", spacing=14)
    d.multiline_text((100, y + 230),
                     "Pagamento: bonifico 60 gg d.f.\n"
                     "IBAN IT60 X054 2811 1010 0000 0123 456\n"
                     "Rif. contratto di fornitura del 30/06/2026",
                     font=mono_s, fill="#3a352c", spacing=10)
    return _effetto_scansione(img, seed, carta="#f6f2e6")


def _rassegna_img(seed: int):
    """Un ritaglio di giornale scansionato, DUE COLONNE: l'OCR semplice
    legge riga per riga e mischia le colonne — serve un tier che capisca
    il layout."""
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1240, 1754), "#efe9d8")
    d = ImageDraw.Draw(img)

    d.text((100, 80), "IL QUOTIDIANO DELL'ARREDO", font=_font_bold(40),
           fill="#171410")
    d.text((100, 140), "sabato 28 giugno 2026 - economia & distretti - pag. 17",
           font=_font_mono(22), fill="#3a352c")
    d.line((100, 190, 1140, 190), fill="#171410", width=3)
    d.multiline_text((100, 230),
                     "L'ufficio che cambia altezza:\nla scommessa di Lumen",
                     font=_font_bold(44), fill="#171410", spacing=10)
    d.line((100, 390, 1140, 390), fill="#171410", width=2)

    mono = _font_mono(25)
    d.multiline_text((100, 430), RASSEGNA_COL1, font=mono, fill="#26221c",
                     spacing=10)
    d.multiline_text((650, 430), RASSEGNA_COL2, font=mono, fill="#26221c",
                     spacing=10)
    d.line((620, 420, 620, 1320), fill="#8d8570", width=1)
    return _effetto_scansione(img, seed, carta="#efe9d8")


def _salva_scansione_pdf(nome: str, immagini, nota: str) -> None:
    import io

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(DATASET / nome), pagesize=A4)
    for img in immagini:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=62)
        buf.seek(0)
        c.drawImage(ImageReader(buf), 0, 0, width=A4[0], height=A4[1])
        c.showPage()
    c.save()
    print(f"  scritto {nome} ({nota})")


def make_verbale_scansione(nome: str) -> None:
    _salva_scansione_pdf(nome, [_pagina_scansionata(VERBALE_PAG1, 41),
                                _pagina_scansionata(VERBALE_PAG2, 42)],
                         "2 pagine solo-immagine, con grana da scanner")


def make_fattura_scansione(nome: str) -> None:
    _salva_scansione_pdf(nome, [_fattura_img(77)],
                         "fattura scansionata CON tabella: OCR + struttura")


def make_rassegna_scansione(nome: str) -> None:
    _salva_scansione_pdf(nome, [_rassegna_img(93)],
                         "ritaglio di giornale a 2 colonne: OCR + layout")


def make_report_corrotto(nome: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = _stili()
    flow = [Paragraph("Report trimestrale Q2 2026 — Gruppo Lumen",
                      styles["Title"])]
    for _ in range(3):                           # ~3 pagine di report plausibile
        for p in REPORT_PARAGRAFI:
            flow.append(Paragraph(p, styles["BodyText"]))
            flow.append(Spacer(1, 7))
    intero = DATASET / "_tmp_intero.pdf"
    hf = _header_footer("Lumen S.r.l. - Report Q2 2026 - riservato")
    SimpleDocTemplate(str(intero), pagesize=A4,
                      topMargin=60, bottomMargin=55).build(
        flow, onFirstPage=hf, onLaterPages=hf)
    dati = intero.read_bytes()
    (DATASET / nome).write_bytes(dati[: int(len(dati) * 0.55)])
    intero.unlink()
    print(f"  scritto {nome} (troncato al 55%: il loader va in errore)")


def make_contratto(nome: str) -> None:
    import docx

    tabelle = {"sla": CONTRATTO_SLA, "corrispettivi": CONTRATTO_CORRISPETTIVI}
    d = docx.Document()
    for tipo, contenuto in CONTRATTO_PAR:
        if tipo == "h1":
            d.add_heading(contenuto, level=1)
        elif tipo == "h2":
            d.add_heading(contenuto, level=2)
        elif tipo == "tabella":
            dati = tabelle[contenuto]
            t = d.add_table(rows=len(dati), cols=len(dati[0]))
            t.style = "Light Grid Accent 1"
            for i, riga in enumerate(dati):
                for j, cella in enumerate(riga):
                    t.rows[i].cells[j].text = cella
        else:
            d.add_paragraph(contenuto)
    d.save(DATASET / nome)
    print(f"  scritto {nome} (8 articoli; le 2 tabelle NON sono nei paragrafi)")


def make_html(nome: str) -> None:
    (DATASET / nome).write_text(FAQ_HTML, encoding="utf-8")
    print(f"  scritto {nome} (10 FAQ + boilerplate e-commerce completo)")


def make_txt_latin1(nome: str) -> None:
    (DATASET / nome).write_bytes(NOTE_INTERNE.encode("latin-1"))
    print(f"  scritto {nome} (encoding latin-1: utf-8 va in errore)")


if __name__ == "__main__":
    DATASET.mkdir(exist_ok=True)
    make_guida("guida_smartworking.pdf")
    make_listino("listino_prezzi.pdf")
    make_verbale_scansione("verbale_scansionato.pdf")
    make_fattura_scansione("fattura_scansionata.pdf")
    make_rassegna_scansione("rassegna_stampa.pdf")
    make_report_corrotto("report_trimestrale.pdf")
    make_html("faq_prodotto.html")
    make_contratto("contratto_fornitura.docx")
    make_txt_latin1("note_interne.txt")
    print("\nDataset pronto in dataset/ — 9 file: trappole diverse, e tier diversi.")
