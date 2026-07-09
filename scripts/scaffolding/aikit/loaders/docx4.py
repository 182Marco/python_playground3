from .base import Document
import docx


def load_docx(path: str) -> list[Document]:
    all = docx.Document(path).paragraphs
    text = "\n".join(p.text for p in all if p.text)
    return [Document(text, {"source": path, "type": "docx"})]