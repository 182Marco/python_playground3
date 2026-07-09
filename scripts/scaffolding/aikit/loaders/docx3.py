import docx
from .base import Document

def load_docx(path: str) -> list[Document]:
    all = docx.Document(path).paragraphs
    txt = "\n".join(p.text for p in all if p.text)
    return [Document(txt, {"source": path, "type": "docx"})]