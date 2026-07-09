"""DOCX → testo dai paragrafi dell'XML (python-docx). Di solito un unico Document.

Un .docx è un archivio zip di XML: il testo c'è, ma annegato nei tag. `python-docx`
ce lo espone come lista di paragrafi.

Nota: `import docx` qui prende il pacchetto installato **python-docx** (import
assoluto). Questo file, per Python, è `loaders.docx` — non c'è collisione.
"""
import docx

from .base import Document


def load_docx(path: str) -> list[Document]:
    documento = docx.Document(path)
    text = "\n".join(p.text for p in documento.paragraphs if p.text)
    return [Document(text, {"source": path, "type": "docx"})]
